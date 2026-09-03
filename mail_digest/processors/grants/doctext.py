"""附件文档 → 文本提取（场景二）。

支持：.docx（python-docx）、.pdf（pypdf）、.xlsx/.xlsm（openpyxl）、
      .csv/.txt/.md（直接读）。
老格式 .doc/.wps/.rtf/.xls 等：若系统装有 LibreOffice（soffice）则转换后读取，
      否则明确报告需人工查看。
rar/.wps 图片等不支持并明确报告。
原则（README）：能读的读，读不了的明确报告，不允许静默失败。
"""
from __future__ import annotations

import csv
import shutil
import subprocess
import tempfile
from pathlib import Path


def _docx_text(path: Path) -> str:
    try:
        import docx
    except ImportError:
        raise ExtractionError("未安装 python-docx")
    d = docx.Document(str(path))
    parts: list[str] = []
    for para in d.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    for table in d.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def _pdf_text(path: Path) -> str:
    try:
        from pypdf import PdfReader
    except ImportError:
        raise ExtractionError("未安装 pypdf")
    reader = PdfReader(str(path))
    pages: list[str] = []
    blank = 0
    for page in reader.pages:
        t = (page.extract_text() or "").strip()
        if t:
            pages.append(t)
        else:
            blank += 1
    note = f"\n[注：{blank}/{len(reader.pages)} 页未提取到文本，可能是扫描件，需 OCR/人工查看]" if blank else ""
    return "\n".join(pages) + note


def _xlsx_text(path: Path) -> str:
    try:
        import openpyxl
    except ImportError:
        raise ExtractionError("未安装 openpyxl")
    wb = openpyxl.load_workbook(str(path), read_only=True, data_only=True)
    parts: list[str] = []
    for ws in wb.worksheets:
        rows: list[str] = []
        for row in ws.iter_rows(values_only=True):
            vals = ["" if v is None else str(v).strip() for v in row]
            if any(vals):
                rows.append(" | ".join(vals))
        if rows:
            parts.append(f"== 工作表: {ws.title} ==")
            parts.extend(rows)
    wb.close()
    return "\n".join(parts)


def _plain_text(path: Path) -> str:
    try:
        return path.read_text(encoding="utf-8", errors="replace")
    except OSError as exc:
        raise ExtractionError(f"读取失败: {exc}")


def _libreoffice_text(path: Path) -> str:
    """老格式（.doc/.wps/.rtf/.xls…）经 LibreOffice headless 转 UTF-8 文本。"""
    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    if not soffice:
        raise ExtractionError(f"{path.suffix} 老格式需要 LibreOffice（soffice），当前未安装，需人工查看")
    with tempfile.TemporaryDirectory() as td:
        try:
            sub = subprocess.run(
                [soffice, "--headless", "--convert-to", "txt:Text (encoded):UTF8",
                 "--outdir", td, str(path)],
                capture_output=True, text=True, timeout=180,
            )
        except (OSError, subprocess.TimeoutExpired) as exc:
            raise ExtractionError(f"LibreOffice 转换失败: {exc}")
        if sub.returncode != 0:
            raise ExtractionError(f"LibreOffice 转换失败: {sub.stderr[:100]}")
        out = Path(td) / (path.stem + ".txt")
        if not out.exists():
            raise ExtractionError("LibreOffice 转换未产出文本（文件可能加密或损坏）")
        text = out.read_text(encoding="utf-8", errors="replace")
        return text.strip() or "（LibreOffice 转换结果为空，可能为扫描/图片型文档）"


class ExtractionError(Exception):
    pass


def extract_text(path: Path) -> str:
    """按扩展名提取文本；不支持/失败抛 ExtractionError（message 可直接展示）。"""
    try:
        return _extract_by_type(path)
    except ExtractionError:
        raise
    except Exception as exc:  # 单附件损坏（如 BadZipFile）只降级为问题报告
        raise ExtractionError(f"解析失败: {exc}")


def _extract_by_type(path: Path) -> str:
    ext = path.suffix.lower()
    if ext == ".docx":
        return _docx_text(path)
    if ext == ".pdf":
        return _pdf_text(path)
    if ext in (".xlsx", ".xlsm"):
        try:
            return _xlsx_text(path)
        except Exception:
            # 可能是老版 .xls 误命名为 .xlsx → 交给 LibreOffice 兜底
            return _libreoffice_text(path)
    if ext in (".csv",):
        # csv 读取（兼容 GBK）
        raw = path.read_bytes()
        for enc in ("utf-8-sig", "gb18030"):
            try:
                text = raw.decode(enc)
                break
            except UnicodeDecodeError:
                continue
        else:
            text = raw.decode("utf-8", errors="replace")
        rows = [" | ".join(r) for r in csv.reader(text.splitlines()) if any(r)]
        return "\n".join(rows)
    if ext in (".txt", ".md", ".log", ".eml"):
        return _plain_text(path)
    if ext in (".doc", ".wps", ".rtf", ".xls", ".ppt", ".pptx"):
        return _libreoffice_text(path)
    if ext in (".rar", ".png", ".jpg", ".jpeg", ".gif", ".bmp", ".exe"):
        raise ExtractionError(f"{ext} 格式当前环境无法解析文本，需人工查看")
    # 未知扩展名：尝试按文本读，失败则报告
    try:
        return _plain_text(path)
    except Exception as exc:
        raise ExtractionError(f"未知格式 {ext}，无法解析: {exc}")


def extract_many(paths: list[Path], max_chars: int = 12000) -> tuple[str, list[dict]]:
    """批量提取：返回 (合并文本, 问题报告列表)。

    合并文本按附件组织，总量超过 max_chars 时截断并在末尾注明。
    """
    parts: list[str] = []
    problems: list[dict] = []
    total = 0
    for p in paths:
        try:
            text = extract_text(p)
        except ExtractionError as exc:
            problems.append({"file": p.name, "error": str(exc)})
            continue
        if not text.strip():
            problems.append({"file": p.name, "error": "文件无可用文本"})
            continue
        head = f"### 附件: {p.name}"
        parts.append(head)
        parts.append(text.strip())
        total += len(text)
    joined = "\n\n".join(parts)
    if total > max_chars:
        joined = joined[:max_chars] + "\n\n[注：附件文本过长，已截断，其余部分略]"
    return joined, problems
